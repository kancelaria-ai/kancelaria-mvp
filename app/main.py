from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from io import BytesIO
import os
from app.analyze_template import analyze_template_with_gpt

app = FastAPI()

# KONKRETNA ŚCIEŻKA do folderu z szablonami
templates = Jinja2Templates(directory="app/templates")

# ========== STRONA GŁÓWNA ==========
@app.get("/", response_class=HTMLResponse)
async def form_page(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# ========== ANALIZA I WYNIK HTML ==========
@app.post("/analyze-template/html", response_class=HTMLResponse)
async def analyze_template_html(request: Request, file: UploadFile = File(...)):
    try:
        contents = await file.read()
        document = Document(BytesIO(contents))
        full_text = "\n".join([para.text for para in document.paragraphs])
        analysis = analyze_template_with_gpt(full_text)

        return templates.TemplateResponse("analysis_result.html", {
            "request": request,
            "contract_type": analysis["contract_type"],
            "summary": analysis["summary"],
            "key_clauses": analysis["key_clauses"],
            "issues_found": analysis["issues_found"],
            "risks": analysis["risks"],
            "recommendations": analysis["recommendations"]
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Błąd analizy: {str(e)}")


# ========== ANALIZA I POBRANIE DOCX ==========
@app.post("/analyze-template-docx")
async def analyze_template_docx(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        doc = Document(BytesIO(contents))
        text = "\n".join(p.text for p in doc.paragraphs)
        analysis = analyze_template_with_gpt(text)

        output_doc = Document()
        output_doc.add_heading("Wynik analizy umowy", level=1)

        output_doc.add_heading("Rodzaj umowy", level=2)
        output_doc.add_paragraph(analysis["contract_type"])

        output_doc.add_heading("Opis", level=2)
        output_doc.add_paragraph(analysis["summary"])

        output_doc.add_heading("Najważniejsze postanowienia", level=2)
        for clause in analysis["key_clauses"]:
            output_doc.add_paragraph(clause, style="List Bullet")

        output_doc.add_heading("Wykryte problemy", level=2)
        for issue in analysis["issues_found"]:
            output_doc.add_paragraph(issue, style="List Bullet")

        output_doc.add_heading("Potencjalne ryzyka", level=2)
        for risk in analysis["risks"]:
            output_doc.add_paragraph(risk, style="List Bullet")

        output_doc.add_heading("Rekomendacje", level=2)
        for rec in analysis["recommendations"]:
            output_doc.add_paragraph(rec, style="List Bullet")

        buf = BytesIO()
        output_doc.save(buf)
        buf.seek(0)

        filename = file.filename.replace(".docx", "_analiza.docx")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Błąd podczas analizy i generowania DOCX: {str(e)}")
